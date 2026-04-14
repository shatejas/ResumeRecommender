"""Create formatted Word documents from structured resume data."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from src.resume_model import ResumeData, DEFAULT_SECTIONS

FONT_NAME = "Times New Roman"
FONT_SIZE_NAME = Pt(14)
FONT_SIZE_HEADING = Pt(10.5)
FONT_SIZE_BODY = Pt(10)
LINE_SPACING = Pt(12)


def _add_run(paragraph, text, bold=False, italic=False, size=FONT_SIZE_BODY):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    return run


def _add_section_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    _add_run(p, title.upper(), bold=True, size=FONT_SIZE_HEADING)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4",
        qn("w:space"): "1", qn("w:color"): "666666",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_summary(doc, data, page_width):
    _add_section_header(doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    _add_run(p, data.summary)


def _render_skills(doc, data, page_width):
    _add_section_header(doc, "Skills")
    for category, items in data.skills.items():
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_run(p, f"{category}: ", bold=True)
        _add_run(p, items)


def _render_experience(doc, data, page_width):
    _add_section_header(doc, "Experience")
    for role in data.experience:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)
        _add_run(p, f"{role.company} | {role.title}", bold=True)
        _add_run(p, "\t")
        _add_run(p, role.dates, italic=True)

        for bullet in role.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.line_spacing = LINE_SPACING
            _add_run(bp, bullet)


def _render_education(doc, data, page_width):
    _add_section_header(doc, "Education")
    for edu in data.education:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)
        label = f"{edu.degree} | {edu.university}"
        if edu.gpa:
            label += f", GPA: {edu.gpa}"
        _add_run(p, label)
        _add_run(p, "\t")
        _add_run(p, edu.date, italic=True)


def _render_certifications(doc, data, page_width):
    if not data.certifications:
        return
    _add_section_header(doc, "Certifications")
    for cert in data.certifications:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_run(p, cert)


def _render_projects(doc, data, page_width):
    if not data.projects:
        return
    _add_section_header(doc, "Projects")
    for proj in data.projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_run(p, proj.name, bold=True)
        if proj.description:
            _add_run(p, f" — {proj.description}")

        for bullet in proj.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.line_spacing = LINE_SPACING
            _add_run(bp, bullet)


def _render_custom(doc, data, section_name, page_width):
    """Render a custom section by name from custom_sections dict."""
    items = data.custom_sections.get(section_name, [])
    if not items:
        return
    _add_section_header(doc, section_name)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_run(p, item)


# Map section names to render functions
_SECTION_RENDERERS = {
    "summary": _render_summary,
    "skills": _render_skills,
    "experience": _render_experience,
    "education": _render_education,
    "certifications": _render_certifications,
    "projects": _render_projects,
}


def save_structured_docx(data: ResumeData, output_path: str,
                         sections: list[str] = None):
    """Create a formatted Word document from structured resume data.
    Sections are rendered in the order specified."""
    if sections is None:
        sections = DEFAULT_SECTIONS

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _add_run(p, data.name, bold=True, size=FONT_SIZE_NAME)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    _add_run(p, data.contact)

    # Render sections in configured order
    for section_name in sections:
        renderer = _SECTION_RENDERERS.get(section_name)
        if renderer:
            renderer(doc, data, page_width)
        else:
            _render_custom(doc, data, section_name, page_width)

    doc.save(output_path)
