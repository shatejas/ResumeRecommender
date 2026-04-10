"""Convert resume text to a formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

FONT_NAME = "Times New Roman"
FONT_SIZE_NAME = Pt(14)
FONT_SIZE_HEADING = Pt(10.5)
FONT_SIZE_BODY = Pt(10)
LINE_SPACING = Pt(12)

SECTION_HEADERS = {
    "summary", "experience", "skills", "education", "certifications",
    "projects", "achievements", "technical skills", "professional experience",
    "education & certifications", "education and certifications",
    "skills & tools", "skills and tools",
    "professional summary", "work experience",
}

SKILLS_HEADERS = {"skills", "technical skills", "skills & tools", "skills and tools"}
EDUCATION_HEADERS = {"education", "education & certifications", "education and certifications"}

EXPERIENCE_LINE_RE = re.compile(
    r"^(.+?\|.+?)\|?\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*[–\-—to]+\s*(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}))\s*$",
    re.IGNORECASE,
)

SKILL_LINE_RE = re.compile(r"^([^:]+):\s*(.+)$")


def _is_section_header(line: str) -> bool:
    cleaned = re.sub(r"[#*:\-]", "", line).strip().lower()
    if len(cleaned) > 40 or cleaned.endswith("."):
        return False
    normalized = re.sub(r"\s*[&]\s*|\s+and\s+", " ", cleaned)
    words = set(normalized.split())
    header_words = {"summary", "experience", "skills", "education", "certifications",
                    "projects", "achievements", "professional", "technical", "work"}
    return bool(words & header_words) and len(words) <= 4


def _get_section_type(line: str) -> str:
    """Return which section type a header belongs to."""
    cleaned = re.sub(r"[#*:\-]", "", line).strip().lower()
    if any(h in cleaned for h in SKILLS_HEADERS):
        return "skills"
    if any(h in cleaned for h in EDUCATION_HEADERS):
        return "education"
    return ""


def _clean_markdown(line: str) -> str:
    line = re.sub(r"^\#{1,4}\s*", "", line)
    line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
    line = re.sub(r"\*(.*?)\*", r"\1", line)
    return line.strip()


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[-•*+]\s+", line))


def _clean_bullet(line: str) -> str:
    return re.sub(r"^\s*[-•*+]\s+", "", line).strip()


def _add_run(paragraph, text, bold=False, italic=False, size=FONT_SIZE_BODY):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    return run


def _add_section_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "666666",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_experience_line(doc, company_title, dates, page_width):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)
    _add_run(p, company_title.strip().rstrip("|").strip(), bold=True)
    _add_run(p, "\t")
    _add_run(p, dates.strip(), italic=True)
    return p


# Matches year like 2019, 2023, etc.
EDU_DATE_RE = re.compile(
    r",?\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+)?\d{4}\s*$",
    re.IGNORECASE,
)


def _add_education_line(doc, text, page_width):
    """Format education: bold degree | university, GPA with right-aligned date."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)

    # Extract date (optional month + year) from end
    date = ""
    date_match = EDU_DATE_RE.search(text)
    if date_match:
        date = date_match.group(0).strip().lstrip(",").strip()
        main_text = text[:date_match.start()].rstrip(", —-–\t ")
    else:
        main_text = text.strip()

    # Split on — or |
    parts = re.split(r"\s*[—|]\s*", main_text, maxsplit=1)

    if len(parts) == 2:
        _add_run(p, parts[0].strip() + " | " + parts[1].strip(), bold=True)
    else:
        _add_run(p, main_text.strip(), bold=True)

    if date:
        _add_run(p, "\t")
        _add_run(p, date, italic=True)

    return p


def _add_skill_line(doc, text):
    """Add a skill line as bullet with bold category."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING

    match = SKILL_LINE_RE.match(text)
    if match:
        _add_run(p, match.group(1).strip() + ": ", bold=True)
        _add_run(p, match.group(2).strip())
    else:
        _add_run(p, text)
    return p


def save_resume_docx(resume_text: str, output_path: str):
    """Parse markdown resume and write a formatted Word document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    lines = resume_text.strip().splitlines()
    content_index = 0
    current_section = ""

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        cleaned = _clean_markdown(stripped)
        if not cleaned:
            continue

        # Name (first non-empty line)
        if content_index == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _add_run(p, cleaned, bold=True, size=FONT_SIZE_NAME)
            content_index += 1
            continue

        # Contact info (second line, usually has | or @ or phone)
        if content_index == 1 and not _is_section_header(stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            _add_run(p, cleaned)
            content_index += 1
            continue

        # Section header
        if _is_section_header(stripped):
            current_section = _get_section_type(stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            _add_run(p, cleaned.upper(), bold=True, size=FONT_SIZE_HEADING)
            _add_section_border(p)
            content_index += 1
            continue

        # Experience line: "Company | Title | dates"
        match = EXPERIENCE_LINE_RE.match(cleaned)
        if match:
            current_section = ""
            _add_experience_line(doc, match.group(1), match.group(2), page_width)
            content_index += 1
            continue

        # Skills section — format as bullet with bold category
        if current_section == "skills" and not _is_bullet(stripped):
            _add_skill_line(doc, cleaned)
            content_index += 1
            continue

        # Education section — bold degree | university, right-aligned year
        if current_section == "education" and not _is_section_header(stripped):
            _add_education_line(doc, cleaned, page_width)
            content_index += 1
            continue

        # Bullet point
        if _is_bullet(stripped):
            bullet_text = _clean_bullet(cleaned)
            if current_section == "skills":
                _add_skill_line(doc, bullet_text)
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = LINE_SPACING
                _add_run(p, bullet_text)
            content_index += 1
            continue

        # Regular text
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_run(p, cleaned)
        content_index += 1

    doc.save(output_path)
